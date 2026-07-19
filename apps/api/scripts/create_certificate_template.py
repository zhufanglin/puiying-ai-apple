"""生成奖状 Word 模板

运行一次即可：
    python scripts/create_certificate_template.py

生成文件：apps/api/templates/apple/certificate.docx
模板中包含 {{ }} 占位符，运行时由 CertificateService 替换。
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docxtpl import DocxTemplate, RichText

TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "templates" / "apple"
OUTPUT_PATH = TEMPLATE_DIR / "certificate.docx"


def create():
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(29.7)   # A4 横向
    section.page_height = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # ---- 标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{ school_name }}")
    run.font.size = Pt(26)
    run.font.bold = True

    # ---- 奖状正文 ----
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.space_before = Pt(20)
    run = body.add_run("兹证明 {{ student_name }} 同学")
    run.font.size = Pt(18)
    body2 = doc.add_paragraph()
    body2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = body2.add_run("在 {{ class_name }} 就读期间，荣获")
    run.font.size = Pt(18)
    body3 = doc.add_paragraph()
    body3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body3.space_after = Pt(10)
    run = body3.add_run("「 {{ award_name }} 」")
    run.font.size = Pt(20)
    run.font.bold = True

    # ---- 备注 ----
    if_remarks = doc.add_paragraph()
    if_remarks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if_remarks.space_before = Pt(10)
    run = if_remarks.add_run("{% if remarks %}特此表彰：{{ remarks }}{% endif %}")
    run.font.size = Pt(14)

    # ---- 日期和签名 ----
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(30)

    footer_left = doc.add_paragraph()
    footer_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_left.add_run("日期：{{ date }}")
    run.font.size = Pt(14)

    footer_right = doc.add_paragraph()
    footer_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_right.add_run("{{ signatory }}")
    run.font.size = Pt(14)

    doc.save(str(OUTPUT_PATH))
    print(f"✅ 模板已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    create()
