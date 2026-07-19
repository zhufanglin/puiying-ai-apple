"""生成正式在学证明 Word 模板。

样式基于 documents skill 的 standard_business_brief，采用命名覆盖
official_certificate_a4：香港学校正式证明使用 A4、居中校名和无边框信头。
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("student_certificate.docx")
INK = RGBColor(29, 41, 57)
MUTED = RGBColor(102, 112, 133)
BRAND = RGBColor(23, 79, 73)


def set_font(run, size: float, *, bold: bool = False, color: RGBColor = INK, latin: str = "Arial") -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (
        ("Heading 1", 16, BRAND, 16, 8),
        ("Heading 2", 13, BRAND, 12, 6),
        ("Heading 3", 12, BRAND, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "174F49")
    borders.append(bottom)
    p_pr.append(borders)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("{{ school_name_zh }}")
    set_font(run, 13, bold=True, color=BRAND)
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("{{ school_name_en }}")
    set_font(run, 9.5, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("{{ title_zh }}")
    set_font(run, 22, bold=True, color=BRAND)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("{{ title_en }}")
    set_font(run, 12.5, bold=True, color=MUTED)
    add_bottom_border(subtitle)

    zh = doc.add_paragraph()
    zh.paragraph_format.space_before = Pt(18)
    zh.paragraph_format.space_after = Pt(16)
    zh.paragraph_format.line_spacing = 1.55
    run = zh.add_run(
        "茲證明本校學生 {{ student_name_zh }}（英文姓名：{{ student_name_en }}；"
        "學號：{{ student_no }}），現就讀 {{ class_name }} 班。該生入學日期為 "
        "{{ admission_date }}。本證明應申請人要求簽發，用途為：{{ purpose }}。"
    )
    set_font(run, 12)

    en = doc.add_paragraph()
    en.paragraph_format.space_after = Pt(28)
    en.paragraph_format.line_spacing = 1.35
    run = en.add_run(
        "This is to certify that {{ student_name_en }} (Chinese name: {{ student_name_zh }}; "
        "Student No. {{ student_no }}) is currently enrolled in Class {{ class_name }} at our school. "
        "The student was admitted on {{ admission_date }}. This certificate is issued upon request "
        "for the following purpose: {{ purpose }}."
    )
    set_font(run, 11)

    date_zh = doc.add_paragraph()
    date_zh.paragraph_format.space_before = Pt(18)
    date_zh.paragraph_format.space_after = Pt(2)
    run = date_zh.add_run("簽發日期：{{ issue_date_zh }}")
    set_font(run, 10.5, color=MUTED)
    date_en = doc.add_paragraph()
    date_en.paragraph_format.space_after = Pt(38)
    run = date_en.add_run("Date of Issue: {{ issue_date_en }}")
    set_font(run, 9.5, color=MUTED)

    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.paragraph_format.space_after = Pt(5)
    run = signature.add_run("校長簽署 / Principal: ____________________")
    set_font(run, 10.5)
    seal = doc.add_paragraph()
    seal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = seal.add_run("學校印章 / School Chop")
    set_font(run, 9.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本證明須經學校簽署及蓋章方為有效。This certificate is valid only with an authorized signature and school chop.")
    set_font(run, 8, color=MUTED)

    doc.core_properties.title = "香港培英中學在學證明模板"
    doc.core_properties.subject = "Apple 学生事务模块"
    doc.core_properties.author = "Hong Kong Pui Ying Secondary School"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
